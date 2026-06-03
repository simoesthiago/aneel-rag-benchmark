"""
test_ingestion.py — Testes unitários da Camada 1 (Ingestão).

Todos os testes rodam OFFLINE — sem bater em APIs reais.
Usamos unittest.mock para simular respostas HTTP.

Para rodar:
    make test
    # ou: pytest tests/test_ingestion.py -v
"""

import pytest
import pandas as pd
from unittest.mock import patch, MagicMock

# =========================================================================
# Testes do pacote extractors/
# =========================================================================


class TestExtrairTextoHtml:
    """Testes da extração de HTML — roda local, sem dependências pesadas."""

    def test_extrai_texto_de_html_simples(self):
        """HTML com parágrafos retorna texto limpo."""
        from src.ingestion.extractors import extrair_texto_html

        html = """
        <html>
        <head><style>body { color: red; }</style></head>
        <body>
            <nav>Menu de navegação</nav>
            <p>Artigo 1º - Esta lei dispõe sobre as concessões de energia.</p>
            <p>Artigo 2º - A ANEEL é responsável pela regulação do setor.</p>
            <footer>Rodapé do site</footer>
        </body>
        </html>
        """
        resultado = extrair_texto_html(html)

        assert "concessões de energia" in resultado["texto"]
        assert "regulação do setor" in resultado["texto"]
        assert "Menu de navegação" not in resultado["texto"]
        assert resultado["extrator"] == "html_parser"
        assert resultado["formato_saida"] == "texto"
        assert resultado["num_paginas"] is None

    def test_html_vazio_retorna_baixa_qualidade(self):
        """HTML sem conteúdo substancial retorna qualidade 0.5."""
        from src.ingestion.extractors import extrair_texto_html

        html = "<html><body><p>Ok</p></body></html>"
        resultado = extrair_texto_html(html)

        assert resultado["qualidade_extracao"] == 0.5

    def test_html_com_conteudo_retorna_alta_qualidade(self):
        """HTML com conteúdo substancial (>1000 chars) retorna qualidade 1.0."""
        from src.ingestion.extractors import extrair_texto_html

        texto_longo = "A" * 500
        html = f"<html><body><p>{texto_longo}</p><p>{texto_longo}</p></body></html>"
        resultado = extrair_texto_html(html)

        assert resultado["qualidade_extracao"] == 1.0


class TestExtrairTextoDispatcher:
    """Testes do dispatcher extrair_texto()."""

    def test_formato_desconhecido_levanta_erro(self):
        """Formato inválido levanta ValueError com mensagem clara."""
        from src.ingestion.extractors import extrair_texto

        with pytest.raises(ValueError, match="não é suportado"):
            extrair_texto("conteudo", "xyz")

    def test_formato_saida_invalido_levanta_erro(self):
        """formato_saida inválido (ex.: "json") levanta ValueError."""
        from src.ingestion.extractors import extrair_texto

        with pytest.raises(ValueError, match="não é válido"):
            extrair_texto(b"bytes", "pdf", formato_saida="json")

    def test_docx_extrai_paragrafos(self):
        """DOCX com parágrafos retorna texto e metodo python_docx."""
        from docx import Document
        from src.ingestion.extractors import extrair_texto
        import io

        doc = Document()
        doc.add_paragraph("Manual de instruções da ANEEL para envio de dados.")
        doc.add_paragraph("Segundo parágrafo com conteúdo regulatório suficiente.")
        buf = io.BytesIO()
        doc.save(buf)

        resultado = extrair_texto(buf.getvalue(), "docx")
        assert "Manual de instruções" in resultado["texto"]
        assert resultado["extrator"] == "python_docx"
        assert resultado["formato_saida"] == "texto"

    def test_xlsx_extrai_celulas(self):
        """XLSX com células retorna texto concatenado."""
        from openpyxl import Workbook
        from src.ingestion.extractors import extrair_texto
        import io

        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Planilha de referência tarifária da distribuidora."
        ws["A2"] = "Dados complementares para o regulador."
        buf = io.BytesIO()
        wb.save(buf)

        resultado = extrair_texto(buf.getvalue(), "xlsx")
        assert "tarifária" in resultado["texto"]
        assert resultado["extrator"] == "openpyxl"
        assert resultado["formato_saida"] == "texto"

    def test_html_aceita_string(self):
        """extrair_texto("html") aceita string como conteúdo."""
        from src.ingestion.extractors import extrair_texto

        resultado = extrair_texto(
            "<html><body><p>Teste de extração de texto HTML.</p></body></html>",
            "html",
        )
        assert "Teste de extração" in resultado["texto"]

    def test_html_aceita_bytes_utf8(self):
        """extrair_texto("html") decodifica bytes UTF-8 automaticamente."""
        from src.ingestion.extractors import extrair_texto

        html_bytes = (
            "<html><body><p>Resolução normativa da ANEEL.</p></body></html>".encode(
                "utf-8"
            )
        )
        resultado = extrair_texto(html_bytes, "html")
        assert "Resolução" in resultado["texto"]

    def test_pdf_exige_bytes(self):
        """extrair_texto("pdf") levanta TypeError se receber string."""
        from src.ingestion.extractors import extrair_texto

        with pytest.raises(TypeError, match="bytes"):
            extrair_texto("string", "pdf")

    def test_docx_markdown_via_mammoth(self):
        """DOCX com formato_saida="markdown" usa mammoth + html2text."""
        from docx import Document
        from src.ingestion.extractors import extrair_texto
        import io

        # python-docx por si só não cria estilos "Heading"; um parágrafo simples
        # já produz Markdown válido (texto puro / parágrafo). O ponto do teste é
        # validar o roteamento: formato_saida=markdown deve cair em mammoth, não
        # em python_docx.
        doc = Document()
        doc.add_heading("Manual de Operação da Distribuidora", level=1)
        doc.add_paragraph(
            "Este documento descreve os procedimentos regulatórios obrigatórios "
            "para as concessionárias de distribuição de energia elétrica."
        )
        buf = io.BytesIO()
        doc.save(buf)

        resultado = extrair_texto(buf.getvalue(), "docx", formato_saida="markdown")
        assert resultado["extrator"] == "mammoth"
        assert resultado["formato_saida"] == "markdown"
        # Heading 1 do python-docx vira "# ..." no Markdown
        assert "Manual de Operação" in resultado["texto"]

    def test_xlsx_markdown_via_pandas(self):
        """XLSX com formato_saida="markdown" gera tabela Markdown por aba."""
        from openpyxl import Workbook
        from src.ingestion.extractors import extrair_texto
        import io

        wb = Workbook()
        ws = wb.active
        ws.title = "Tarifas"
        ws.append(["distribuidora", "tarifa_kwh"])
        ws.append(["CPFL", 0.65])
        ws.append(["Light", 0.72])
        buf = io.BytesIO()
        wb.save(buf)

        resultado = extrair_texto(buf.getvalue(), "xlsx", formato_saida="markdown")
        assert resultado["extrator"] == "pandas_tabulate"
        assert resultado["formato_saida"] == "markdown"
        assert "## Tarifas" in resultado["texto"]
        # df.to_markdown() usa "|" como separador de colunas
        assert "|" in resultado["texto"]
        assert "distribuidora" in resultado["texto"]


# =========================================================================
# Testes do scraper_atos.py
# =========================================================================


class TestMontarUrlAto:
    """Testes da construção de URLs do cedoc/."""

    def test_url_padrao_ren(self):
        """REN 1000/2021 → ren20211000.pdf"""
        from src.ingestion.scraper_atos import montar_url_ato

        url = montar_url_ato("ren", 2021, 1000)
        assert url == "https://www2.aneel.gov.br/cedoc/ren20211000.pdf"

    def test_url_sem_zero_pad(self):
        """REN 414/2010 → ren2010414.pdf (sem zero-padding, confirmado no cedoc/)"""
        from src.ingestion.scraper_atos import montar_url_ato

        url = montar_url_ato("ren", 2010, 414)
        assert url == "https://www2.aneel.gov.br/cedoc/ren2010414.pdf"

    def test_url_consolidada(self):
        """Versão consolidada tem prefixo 'b'."""
        from src.ingestion.scraper_atos import montar_url_consolidada

        url = montar_url_consolidada("ren", 2021, 1000)
        assert url == "https://www2.aneel.gov.br/cedoc/bren20211000.pdf"

    def test_url_res(self):
        """Funciona para outros tipos de ato (RES), também sem zero-padding."""
        from src.ingestion.scraper_atos import montar_url_ato

        url = montar_url_ato("res", 2002, 798)
        assert url == "https://www2.aneel.gov.br/cedoc/res2002798.pdf"


class TestFiltrarNaoVigentes:
    """Testes do filtro de atos não vigentes."""

    def test_filtra_revogados(self):
        from src.ingestion.scraper_atos import filtrar_nao_vigentes

        atos = [
            {"Resolução": "REN 1000/2021", "Situação": "Não consta revogação expressa"},
            {"Resolução": "REN 200/2005", "Situação": "Revogada"},
        ]
        resultado = filtrar_nao_vigentes(atos, sigla="ren")
        assert len(resultado) == 1
        assert resultado[0]["Resolução"] == "REN 200/2005"


class TestFiltrarVigentes:
    """Testes do filtro de atos vigentes."""

    def test_filtra_por_situacao(self):
        """Apenas atos vigentes passam o filtro."""
        from src.ingestion.scraper_atos import filtrar_vigentes

        atos = [
            {"Resolução": "REN 1000/2021", "Situação": "Não consta revogação expressa"},
            {"Resolução": "REN 200/2005", "Situação": "Revogada"},
            {"Resolução": "REN 500/2015", "Situação": "Vacatio Legis"},
        ]
        resultado = filtrar_vigentes(atos)
        assert len(resultado) == 2
        nomes = [r["Resolução"] for r in resultado]
        assert "REN 1000/2021" in nomes
        assert "REN 500/2015" in nomes

    def test_filtra_por_sigla(self):
        """Filtro por sigla funciona."""
        from src.ingestion.scraper_atos import filtrar_vigentes

        atos = [
            {"Resolução": "REN 1000/2021", "Situação": "Não consta revogação expressa"},
            {"Resolução": "RES 798/2002", "Situação": "Não consta revogação expressa"},
        ]
        resultado = filtrar_vigentes(atos, sigla="ren")
        assert len(resultado) == 1
        assert resultado[0]["Resolução"] == "REN 1000/2021"


class TestConsultarPowerbiMock:
    """Testa o parsing da resposta do Power BI com mock."""

    @patch("src.ingestion.scraper_atos.requests.post")
    def test_decodifica_resposta_basica(self, mock_post):
        """Resposta simples do Power BI é decodificada corretamente."""
        from src.ingestion.scraper_atos import consultar_powerbi

        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.json.return_value = {
            "results": [
                {
                    "result": {
                        "data": {
                            "dsr": {
                                "Version": 2,
                                "DS": [
                                    {
                                        "PH": [
                                            {
                                                "DM0": [
                                                    {
                                                        "S": [
                                                            {
                                                                "N": "G0",
                                                                "T": 1,
                                                                "DN": "D0",
                                                            },
                                                        ],
                                                        "C": [0],
                                                    },
                                                    {"C": [1], "R": 0},
                                                ]
                                            }
                                        ],
                                        "ValueDicts": {
                                            "D0": ["REN 1000/2021", "RES 798/2002"],
                                        },
                                    }
                                ],
                            }
                        }
                    }
                }
            ]
        }
        mock_post.return_value = mock_response

        resultado = consultar_powerbi(["Resolução"])

        assert len(resultado) == 2
        assert resultado[0]["Resolução"] == "REN 1000/2021"
        assert resultado[1]["Resolução"] == "RES 798/2002"


# =========================================================================
# Testes do scraper_leis.py
# =========================================================================


class TestColetarLeisMock:
    """Testa a coleta de leis com mock HTTP."""

    @patch("src.ingestion.scraper_leis.requests.get")
    def test_coleta_lei_simples(self, mock_get):
        """Uma lei com HTML simples é coletada corretamente."""
        from src.ingestion.scraper_leis import coletar_leis

        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.encoding = "utf-8"
        mock_response.text = (
            "<html><body>"
            "<p>Art. 1º Fica instituída a Agência Nacional de Energia Elétrica - ANEEL, "
            "autarquia sob regime especial, vinculada ao Ministério de Minas e Energia, "
            "com sede e foro no Distrito Federal e prazo de duração indeterminado.</p>"
            "<p>Art. 2º A ANEEL tem por finalidade regular e fiscalizar a produção, "
            "transmissão, distribuição e comercialização de energia elétrica, em "
            "conformidade com as políticas e diretrizes do governo federal.</p>"
            "</body></html>"
        )
        mock_get.return_value = mock_response

        documentos = coletar_leis()

        assert len(documentos) == 4
        for doc in documentos:
            assert doc["tipo"] == "lei"
            assert doc["subtipo"] == "lei_federal"
            assert doc["fonte"] == "planalto"
            assert doc["formato_original"] == "html"
            assert "ANEEL" in doc["texto_bruto"]


# =========================================================================
# Testes do schema.py
# =========================================================================


class TestValidarSchema:
    """Testes de validação do schema do DataFrame."""

    def _criar_df_valido(
        self, metodo: str = "texto", extrator: str = "html_parser"
    ) -> pd.DataFrame:
        """Cria um DataFrame mínimo válido para testes."""
        return pd.DataFrame(
            [
                {
                    "id": "lei-9427-1996",
                    "tipo": "lei",
                    "subtipo": "lei_federal",
                    "numero": "9427",
                    "ano": 1996,
                    "titulo": "Lei 9.427/1996 — Criação da ANEEL",
                    "assunto": None,
                    "situacao": None,
                    "data_publicacao": None,
                    "fonte": "planalto",
                    "url_original": "https://www.planalto.gov.br/ccivil_03/leis/l9427cons.htm",
                    "url_consolidado": None,
                    "formato_original": "html",
                    "texto_bruto": "A" * 200,
                    "num_paginas": None,
                    "metodo_extracao": metodo,
                    "extrator": extrator,
                    "qualidade_extracao": 1.0,
                    "hf_path": None,
                    "scraped_at": "2026-05-29T00:00:00Z",
                }
            ]
        )

    def test_df_valido_nao_levanta_erro(self):
        """DataFrame válido passa sem erro."""
        from src.ingestion.schema import validar_schema

        df = self._criar_df_valido()
        validar_schema(df)

    def test_coluna_faltando_levanta_erro(self):
        """DataFrame sem coluna obrigatória levanta RuntimeError."""
        from src.ingestion.schema import validar_schema

        df = self._criar_df_valido()
        df = df.drop(columns=["titulo"])

        with pytest.raises(RuntimeError, match="Colunas faltando"):
            validar_schema(df)

    def test_par_id_metodo_duplicado_levanta_erro(self):
        """Par (id, metodo_extracao) duplicado levanta RuntimeError."""
        from src.ingestion.schema import validar_schema

        df = self._criar_df_valido()
        df = pd.concat([df, df], ignore_index=True)

        with pytest.raises(RuntimeError, match="duplicados"):
            validar_schema(df)

    def test_mesmo_id_estrategias_diferentes_e_valido(self):
        """Mesmo id com metodo_extracao diferente é válido (Opção A)."""
        from src.ingestion.schema import validar_schema

        df_texto = self._criar_df_valido(metodo="texto", extrator="html_parser")
        df_md = self._criar_df_valido(metodo="markdown", extrator="html2markdown")
        df = pd.concat([df_texto, df_md], ignore_index=True)

        validar_schema(df)  # não deve levantar exceção

    def test_metodo_extracao_invalido_levanta_erro(self):
        """metodo_extracao com nome de ferramenta (ex.: "pymupdf") é rejeitado."""
        from src.ingestion.schema import validar_schema

        df = self._criar_df_valido(metodo="pymupdf", extrator="pymupdf")
        with pytest.raises(RuntimeError, match="metodo_extracao com valores inválidos"):
            validar_schema(df)

    def test_texto_vazio_levanta_erro(self):
        """Documento com texto_bruto muito curto levanta RuntimeError."""
        from src.ingestion.schema import validar_schema

        df = self._criar_df_valido()
        df.loc[0, "texto_bruto"] = "curto"

        with pytest.raises(RuntimeError, match="texto_bruto < 100"):
            validar_schema(df)


class TestMesclarCorpus:
    """Testes de merge incremental."""

    def _doc(
        self,
        doc_id: str,
        scraped_at: str,
        metodo: str = "texto",
        extrator: str = "html_parser",
    ) -> dict:
        return {
            "id": doc_id,
            "tipo": "lei",
            "subtipo": "lei_federal",
            "numero": "9427",
            "ano": 1996,
            "titulo": "Lei teste",
            "assunto": None,
            "situacao": None,
            "data_publicacao": None,
            "fonte": "planalto",
            "url_original": "https://example.com",
            "url_consolidado": None,
            "formato_original": "html",
            "texto_bruto": "A" * 200,
            "num_paginas": None,
            "metodo_extracao": metodo,
            "extrator": extrator,
            "qualidade_extracao": 1.0,
            "hf_path": None,
            "scraped_at": scraped_at,
        }

    def test_merge_sem_duplicatas(self):
        from src.ingestion.uploader import mesclar_corpus

        df1 = pd.DataFrame([self._doc("lei-9427-1996", "2026-01-01T00:00:00Z")])
        df2 = pd.DataFrame([self._doc("ren-2020-100", "2026-02-01T00:00:00Z")])
        merged = mesclar_corpus(df1, df2)
        assert len(merged) == 2

    def test_merge_id_conflito_mantem_mais_recente(self):
        from src.ingestion.uploader import mesclar_corpus

        df1 = pd.DataFrame([self._doc("lei-9427-1996", "2026-01-01T00:00:00Z")])
        novo = self._doc("lei-9427-1996", "2026-06-01T12:00:00Z")
        novo["texto_bruto"] = "B" * 200
        df2 = pd.DataFrame([novo])
        merged = mesclar_corpus(df1, df2)
        assert len(merged) == 1
        assert merged.iloc[0]["texto_bruto"].startswith("B")

    def test_merge_mesmo_id_estrategias_diferentes_preserva_ambos(self):
        """Doc com texto + markdown devem coexistir após merge."""
        from src.ingestion.uploader import mesclar_corpus

        df1 = pd.DataFrame(
            [self._doc("ren-2021-1000", "2026-01-01T00:00:00Z", metodo="texto", extrator="pymupdf")]
        )
        df2 = pd.DataFrame(
            [self._doc("ren-2021-1000", "2026-01-01T00:00:00Z", metodo="markdown", extrator="pymupdf4llm")]
        )
        merged = mesclar_corpus(df1, df2)
        assert len(merged) == 2
        assert set(merged["metodo_extracao"]) == {"texto", "markdown"}


class TestColetarProretMock:
    @patch("src.ingestion.scraper_procedimentos.baixar_arquivo_gitlab")
    @patch("src.ingestion.scraper_procedimentos._listar_pdfs_recursivo")
    @patch("src.ingestion.scraper_procedimentos._encontrar_pasta_proret")
    def test_coletar_proret_um_pdf(self, mock_path, mock_listar, mock_baixar):
        from src.ingestion.scraper_procedimentos import coletar_proret

        mock_path.return_value = "procreg/proret"
        mock_listar.return_value = ["procreg/proret/modulo02/subm2.1/doc.pdf"]
        mock_baixar.return_value = b"%PDF-1.4 fake"

        with patch("src.ingestion.scraper_procedimentos.extrair_texto") as mock_ext:
            mock_ext.return_value = {
                "texto": "X" * 150,
                "num_paginas": 1,
                "qualidade_extracao": 1.0,
                "formato_saida": "texto",
                "extrator": "pymupdf",
            }
            docs = coletar_proret()
        assert len(docs) == 1
        assert docs[0]["subtipo"] == "proret"


class TestColetarManuaisMock:
    @patch("src.ingestion.scraper_manuais._baixar_url")
    @patch("src.ingestion.scraper_manuais.requests.get")
    def test_coleta_link_pdf_na_pagina(self, mock_get, mock_baixar):
        from src.ingestion.scraper_manuais import coletar_manuais

        html = """
        <html><body>
        <a href="https://git.aneel.gov.br/publico/x/manual.pdf">Manual teste</a>
        </body></html>
        """
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.text = html
        mock_resp.raise_for_status = MagicMock()
        mock_get.return_value = mock_resp
        mock_baixar.return_value = b"%PDF fake"

        with patch("src.ingestion.scraper_manuais.extrair_texto") as mock_ext:
            mock_ext.return_value = {
                "texto": "Y" * 150,
                "num_paginas": 1,
                "qualidade_extracao": 1.0,
                "formato_saida": "texto",
                "extrator": "pymupdf",
            }
            docs = coletar_manuais(subcategorias=["distribuicao"], max_manuais=5)
        assert len(docs) == 1
        assert docs[0]["tipo"] == "manual"
